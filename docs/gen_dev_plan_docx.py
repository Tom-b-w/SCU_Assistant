"""
生成《SCU Assistant 项目开发计划书》Word 文档
依赖: python-docx (pip install python-docx)
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "项目开发计划书_v3.docx")

# ---------------------------------------------------------------------------
# Color constants
# ---------------------------------------------------------------------------
DARK_BLUE = RGBColor(47, 85, 151)
HEADER_GRAY = "D9E2F3"  # table header bg
GANTT_BLUE = "2F5597"
GANTT_RED = "C00000"
GREEN_BG = "70AD47"
ORANGE_BG = "ED7D31"
PURPLE_BG = "7030A0"
GRAY_BG = "A5A5A5"
LIGHT_GRAY_BG = "D0D0D0"
WHITE_BG = "FFFFFF"
LIGHT_BLUE_BG = "B4C6E7"

# ---------------------------------------------------------------------------
# Helper functions
# ---------------------------------------------------------------------------

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def set_run_font(run, font_name="微软雅黑", size=None, bold=False, color=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def add_styled_paragraph(doc, text, font_name="微软雅黑", size=10.5, bold=False,
                          alignment=None, space_before=0, space_after=6, color=None,
                          first_line_indent=None):
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    run = p.add_run(text)
    set_run_font(run, font_name, size, bold, color)
    return p


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    if level == 1:
        set_run_font(run, "微软雅黑", 18, True, DARK_BLUE)
    elif level == 2:
        set_run_font(run, "微软雅黑", 15, True, DARK_BLUE)
    elif level == 3:
        set_run_font(run, "微软雅黑", 13, True, DARK_BLUE)
    elif level == 4:
        set_run_font(run, "微软雅黑", 11, True, RGBColor(0x40, 0x40, 0x40))
    return h


def create_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        set_run_font(run, "微软雅黑", 10, True, RGBColor(0, 0, 0))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, HEADER_GRAY)

    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            set_run_font(run, "微软雅黑", 9.5)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def add_bullet_list(doc, items, level=0, font_name="微软雅黑", size=10.5):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.clear()
        run = p.add_run(item)
        set_run_font(run, font_name, size)
        if level > 0:
            p.paragraph_format.left_indent = Cm(1.27 * level)


def add_numbered_item(doc, number, title, sub_items):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{number}. {title}")
    set_run_font(run, "微软雅黑", 10.5, True)

    for item in sub_items:
        bp = doc.add_paragraph()
        bp.paragraph_format.left_indent = Cm(1.27)
        bp.paragraph_format.space_before = Pt(0)
        bp.paragraph_format.space_after = Pt(2)
        run = bp.add_run(f"- {item}")
        set_run_font(run, "微软雅黑", 10)


# ---------------------------------------------------------------------------
# Gantt chart helpers
# ---------------------------------------------------------------------------

def create_gantt_table(doc):
    weeks = [f"W{i}" for i in range(1, 13)]
    sprint_labels = [
        "阶段一 需求分析与框架搭建",
        "阶段二 核心功能开发",
        "阶段三 AI接入与功能完善",
        "阶段四 测试优化与交付",
        "里程碑",
    ]

    sprint_weeks = {
        0: list(range(0, 3)),      # W1-W3
        1: list(range(3, 6)),      # W4-W6
        2: list(range(6, 9)),      # W7-W9
        3: list(range(9, 12)),     # W10-W12
    }

    milestones = {2: "M1", 5: "M2", 8: "M3", 11: "M4"}

    table = doc.add_table(rows=len(sprint_labels) + 1, cols=len(weeks) + 1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    # Header row
    hdr_cell = table.rows[0].cells[0]
    hdr_cell.text = ""
    p = hdr_cell.paragraphs[0]
    run = p.add_run("阶段 \\ 周次")
    set_run_font(run, "微软雅黑", 8, True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(hdr_cell, HEADER_GRAY)
    hdr_cell.width = Cm(5.0)

    for i, w in enumerate(weeks):
        cell = table.rows[0].cells[i + 1]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(w)
        set_run_font(run, "微软雅黑", 7.5, True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, HEADER_GRAY)
        cell.width = Cm(1.15)

    for r_idx, label in enumerate(sprint_labels):
        row = table.rows[r_idx + 1]
        lbl_cell = row.cells[0]
        lbl_cell.text = ""
        p = lbl_cell.paragraphs[0]
        run = p.add_run(label)
        set_run_font(run, "微软雅黑", 8, True)
        lbl_cell.width = Cm(5.0)

        for w_idx in range(12):
            cell = row.cells[w_idx + 1]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.width = Cm(1.15)

            if r_idx < 4:
                if w_idx in sprint_weeks.get(r_idx, []):
                    set_cell_shading(cell, GANTT_BLUE)
                    run = p.add_run("\u2588")
                    set_run_font(run, "微软雅黑", 7, False, RGBColor(47, 85, 151))
            elif r_idx == 4:
                if w_idx in milestones:
                    set_cell_shading(cell, GANTT_RED)
                    run = p.add_run(milestones[w_idx])
                    set_run_font(run, "微软雅黑", 7, True, RGBColor(255, 255, 255))

    doc.add_paragraph()
    return table


def create_role_gantt_table(doc):
    weeks = [f"W{i}" for i in range(1, 13)]
    roles = [
        "Team A 学业核心（覃泽锴、孔垂骄）",
        "Team B 生活服务（李亚飞、谭旭睿）",
        "Team C 校园信息（徐锐学、毛立业）",
        "AI 组（张傲楚、谭博文）",
        "运维（朱圣相）",
    ]

    role_config = {
        0: {"color": GANTT_BLUE, "weeks": list(range(0, 12))},
        1: {"color": GREEN_BG,   "weeks": list(range(0, 12))},
        2: {"color": ORANGE_BG,  "weeks": list(range(0, 12))},
        3: {"color": PURPLE_BG,  "weeks": list(range(2, 12))},
        4: {"color": GRAY_BG,    "weeks": [0, 1, 2, 5, 8, 9, 10, 11],
            "heavy": [9, 10, 11], "light": [0, 1, 2]},
    }

    table = doc.add_table(rows=len(roles) + 1, cols=len(weeks) + 1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    hdr_cell = table.rows[0].cells[0]
    hdr_cell.text = ""
    p = hdr_cell.paragraphs[0]
    run = p.add_run("小组 \\ 周次")
    set_run_font(run, "微软雅黑", 8, True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(hdr_cell, HEADER_GRAY)
    hdr_cell.width = Cm(5.0)

    for i, w in enumerate(weeks):
        cell = table.rows[0].cells[i + 1]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(w)
        set_run_font(run, "微软雅黑", 7.5, True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, HEADER_GRAY)
        cell.width = Cm(1.15)

    for r_idx, role in enumerate(roles):
        row = table.rows[r_idx + 1]
        lbl_cell = row.cells[0]
        lbl_cell.text = ""
        p = lbl_cell.paragraphs[0]
        run = p.add_run(role)
        set_run_font(run, "微软雅黑", 8, True)
        lbl_cell.width = Cm(5.0)

        cfg = role_config[r_idx]
        for w_idx in range(12):
            cell = row.cells[w_idx + 1]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.width = Cm(1.15)

            if w_idx in cfg["weeks"]:
                color = cfg["color"]
                if r_idx == 4 and w_idx in cfg.get("light", []):
                    color = LIGHT_GRAY_BG
                set_cell_shading(cell, color)
                run = p.add_run("\u2588")
                set_run_font(run, "微软雅黑", 7, False, RGBColor(255, 255, 255))

    doc.add_paragraph()
    return table


# ---------------------------------------------------------------------------
# Main document generation
# ---------------------------------------------------------------------------

def generate_document():
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # -- Default style --
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.paragraph_format.line_spacing = 1.5

    # ===================================================================
    # TITLE PAGE
    # ===================================================================
    for _ in range(4):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(12)
    run = title_p.add_run("SCU Assistant 项目开发计划书")
    set_run_font(run, "微软雅黑", 28, True, DARK_BLUE)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(8)
    run = subtitle_p.add_run("四川大学智能校园助手")
    set_run_font(run, "微软雅黑", 18, False, RGBColor(0x66, 0x66, 0x66))

    course_p = doc.add_paragraph()
    course_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    course_p.paragraph_format.space_after = Pt(24)
    run = course_p.add_run("软件工程综合实践")
    set_run_font(run, "微软雅黑", 14, False, RGBColor(0x88, 0x88, 0x88))

    team_p = doc.add_paragraph()
    team_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    team_p.paragraph_format.space_after = Pt(6)
    run = team_p.add_run("团队成员：谭博文  徐锐学  孔垂骄  张傲楚  毛立业  朱圣相  李亚飞  谭旭睿  覃泽锴")
    set_run_font(run, "微软雅黑", 11, False, RGBColor(0x55, 0x55, 0x55))

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run("2026 年 3 月 21 日")
    set_run_font(run, "微软雅黑", 12, False, RGBColor(0x99, 0x99, 0x99))

    doc.add_page_break()

    # ===================================================================
    # 1 项目概述
    # ===================================================================
    add_heading_styled(doc, "1  项目概述", level=1)

    # 1.1
    add_heading_styled(doc, "1.1  项目背景", level=2)
    add_styled_paragraph(
        doc,
        "SCU Assistant（四川大学智能助手）是面向四川大学在校学生的一站式 AI 校园生活助手。"
        "项目旨在将课业管理、校园服务、生活决策等高频需求整合到统一平台，通过大语言模型（LLM）"
        '提供自然语言交互体验，让学生\u201c问一句话就能办一件事\u201d。',
        first_line_indent=0.74,
    )
    add_styled_paragraph(
        doc,
        "当前四川大学校园信息化存在三大痛点：（1）信息碎片化——课表、成绩、校车、食堂分布在多个独立平台，"
        "学生需在多个系统间切换；（2）系统体验差——教务系统 UI 陈旧，移动端适配差，查课表需多步操作；"
        "（3）缺乏智能化——无 AI 问答入口，无个性化推荐，无统一的信息聚合能力。"
        "四川大学 6 万+在校师生急需一个一站式智能校园助手。",
        first_line_indent=0.74,
    )

    # 1.2
    add_heading_styled(doc, "1.2  当前进展", level=2)
    add_styled_paragraph(
        doc,
        "项目目前处于初步规划阶段（约 10%），已完成选题论证、团队组建、技术选型和功能模块规划，"
        "正在进行技术栈学习和基础框架搭建。具体进展如下：",
        first_line_indent=0.74,
    )

    create_table(doc,
        ["类别", "内容", "状态"],
        [
            ["选题论证", "确定项目选题与整体方案，明确产品定位", "已完成 ✓"],
            ["团队组建", "9 人团队、4 个专业小组（前端/后端/AI/运维）分工明确", "已完成 ✓"],
            ["技术选型", "确定 Next.js + FastAPI + PostgreSQL + Redis 技术栈", "已完成 ✓"],
            ["功能规划", "5 大模块 23 个功能点全景规划完成", "已完成 ✓"],
            ["代码仓库", "GitHub 仓库已创建（github.com/Tom-b-w/SCU_Assistant）", "已完成 ✓"],
            ["前端学习", "团队成员学习 Next.js / React / TypeScript", "进行中 ○"],
            ["后端学习", "团队成员学习 FastAPI / Python / SQLAlchemy", "进行中 ○"],
            ["AI 调研", "学习 AI 大模型接口调用方式、DeepTutor 源码调研", "进行中 ○"],
            ["框架搭建", "前后端项目基础框架初始化", "待开始 ○"],
            ["数据库设计", "数据库表结构设计", "待开始 ○"],
            ["教务调研", "教务系统登录流程与 API 调研", "待开始 ○"],
        ],
        col_widths=[3, 9, 3.5],
    )

    add_styled_paragraph(
        doc,
        "拟采用技术栈：Next.js + React + TypeScript（前端）| FastAPI + SQLAlchemy + PostgreSQL + Redis（后端）| Docker Compose 部署",
        size=9.5, bold=True,
    )

    # 1.3
    add_heading_styled(doc, "1.3  目标功能全景", level=2)
    add_styled_paragraph(
        doc,
        "基于需求分析，项目规划 5 大模块共 23 个功能点。以下为完整功能规划：",
        first_line_indent=0.74,
    )

    # 学业模块
    add_heading_styled(doc, "学业模块（6 个功能）", level=3)
    create_table(doc,
        ["序号", "功能", "说明", "优先级"],
        [
            ["1", "课表智能问答", '自然语言问课表：\u201c明天有什么课\uff1f\u201d\u201c周三下午哪里有空教室\uff1f\u201d', "P0"],
            ["2", "DDL 多源追踪", "手动添加 + 教务通知自动抓取 + 微信/QQ/学习通 DDL 同步，每日推送", "P0"],
            ["3", "选课推荐", "根据培养方案 + 已修课程，算出还差什么，推荐下学期选什么，自动检测时间冲突", "P1"],
            ["4", "考试倒计时 + AI 复习计划", "自动抓考试安排生成倒计时，LLM 生成逐日复习计划", "P1"],
            ["5", "课件 RAG 问答（DeepTutor）", "上传 PPT/PDF 向量化，双循环推理 + RAG + 代码执行，考前针对性提问", "P1"],
            ["6", "一键评教", "期末批量填写教学评价", "P2"],
        ],
        col_widths=[1.2, 4, 8, 2],
    )

    add_styled_paragraph(doc, "DeepTutor 集成说明", size=11, bold=True, space_before=6)
    add_styled_paragraph(
        doc,
        "参考 HKUDS/DeepTutor（https://github.com/HKUDS/DeepTutor）的双循环推理架构，"
        "计划实现以下能力：",
        first_line_indent=0.74,
    )
    add_bullet_list(doc, [
        "文档知识库 Q&A：上传教材/课件建立向量知识库，带精确引用的问答",
        "智能出题：根据上传的参考试卷，生成风格/格式/难度匹配的练习题",
        "知识图谱可视化：自动构建课程知识点关系图，辅助复习",
        "上下文感知对话：根据学习进度自适应调整回答深度",
    ])

    # 吃喝模块
    add_heading_styled(doc, "吃喝模块（5 个功能）", level=3)
    create_table(doc,
        ["序号", "功能", "说明", "优先级"],
        [
            ["1", "食堂营业状态", '三校区所有食堂实时查\u201c哪个食堂还开着\u201d', "P0"],
            ["2", "食堂窗口导览", '各食堂窗口品类标签，\u201c江安哪里有麻辣烫\uff1f\u201d', "P0"],
            ["3", '\u201c今天吃什么\u201d决策器', "品类随机推荐 + 用户口味偏好学习", "P1"],
            ["4", "校外美食推荐", '高德 POI + 美团/饿了么外卖数据，\u201c望江附近好吃的川菜\uff1f\u201d', "P1"],
            ["5", "一卡通余额/消费", '\u201c卡里还有多少钱\uff1f\u201d\u201c这周花了多少\uff1f\u201d', "P2"],
        ],
        col_widths=[1.2, 4, 8, 2],
    )

    # 玩乐模块
    add_heading_styled(doc, "玩乐模块（4 个功能）", level=3)
    create_table(doc,
        ["序号", "功能", "说明", "优先级"],
        [
            ["1", "周末规划", "结合高德 POI + 天气 + LLM 生成行程", "P1"],
            ["2", "校园活动", "自动爬取团委/学院活动通知，AI 摘要", "P1"],
            ["3", "天气 + 穿衣建议", '接入和风天气 API，\u201c明天要带伞吗\uff1f\u201d\u201c穿什么合适\uff1f\u201d', "P0"],
            ["4", "快递追踪", "接入菜鸟/京东/淘宝物流数据，到校园取件点时推送提醒", "P2"],
        ],
        col_widths=[1.2, 4, 8, 2],
    )

    # 校园基础模块
    add_heading_styled(doc, "校园基础模块（4 个功能）", level=3)
    create_table(doc,
        ["序号", "功能", "说明", "优先级"],
        [
            ["1", "图书馆", "查书、借阅记录、座位情况", "P1"],
            ["2", "校车时刻", '望江\u2194江安\u2194华西校车时间表，\u201c下一班去江安的校车几点\uff1f\u201d', "P0"],
            ["3", "校园通知聚合", "教务处 + 软件学院 + 研究生院通知自动抓取，AI 摘要推送", "P1"],
            ["4", "校历查询", '\u201c寒假什么时候开始\uff1f\u201d\u201c现在第几周\uff1f\u201d', "P0"],
        ],
        col_widths=[1.2, 4, 8, 2],
    )

    # Agent 核心层
    add_heading_styled(doc, "Agent 核心层（4 个功能）", level=3)
    create_table(doc,
        ["序号", "功能", "说明", "优先级"],
        [
            ["1", "Morning Briefing", "每天 7:30 推送今日课表 + 天气 + DDL + 考试倒计时 + 食堂状态", "P1"],
            ["2", "自然语言路由", "任意问题自动识别意图，路由到正确的功能模块", "P0"],
            ["3", "跨源推理", '\u201c下雨 + 没课 + 图书馆快满了\u201d\u2192 综合建议\u201c建议现在去占座\u201d', "P2"],
            ["4", "用户记忆", "记住口味偏好、校区、专业，越用越懂你", "P1"],
        ],
        col_widths=[1.2, 4, 8, 2],
    )

    # 数据来源说明
    add_heading_styled(doc, "1.4  数据来源", level=2)
    create_table(doc,
        ["数据源", "获取内容", "接入方式"],
        [
            ["四川大学教务系统", "课表、成绩、学分、GPA、考试安排", "CAS 统一认证模拟登录"],
            ["校园公开信息", "食堂窗口信息、校车时刻表（望江/江安/华西）", "静态数据内置 + 定期更新"],
            ["用户自主录入", "DDL 截止日期、个人待办、口味偏好", "系统内 CRUD 操作"],
            ["AI 大模型", "自然语言问答、意图识别、智能推荐", "通义千问 / DeepSeek REST API"],
            ["微信/QQ/学习通", "课程 DDL、作业通知", "OAuth 授权 / 消息 WebHook / 数据抓取"],
            ["美团/饿了么", "校外餐饮商户、菜品、评分", "开放 API / 数据抓取"],
            ["和风天气 API", "实时天气、7 天预报、生活指数", "REST API（免费 1000 次/天）"],
            ["高德地图 API", "POI 搜索、路线规划", "REST API（免费配额）"],
            ["快递 100 API", "物流跟踪（菜鸟/京东/淘宝聚合）", "REST API"],
            ["四川大学图书馆", "图书检索、借阅记录", "HTTP 爬虫"],
        ],
        col_widths=[3.5, 5.5, 7],
    )

    # ===================================================================
    # 2 实施计划
    # ===================================================================
    add_heading_styled(doc, "2  实施计划", level=1)

    # 2.1
    add_heading_styled(doc, "2.1  开发模型", level=2)
    add_styled_paragraph(
        doc,
        "采用增量迭代 + 敏捷开发模型，共划分 4 个开发阶段，每阶段 3 周，总周期 12 周。"
        "采用 GitHub Flow 协作模式，每周例会对齐进度，双周迭代交付，接口文档先行。",
        first_line_indent=0.74,
    )

    flow_p = doc.add_paragraph()
    flow_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    flow_p.paragraph_format.space_before = Pt(6)
    flow_p.paragraph_format.space_after = Pt(12)
    run = flow_p.add_run(
        "阶段一（需求与搭建） → 阶段二（核心功能） → 阶段三（AI 与完善） → 阶段四（测试与交付）"
    )
    set_run_font(run, "微软雅黑", 10, True, DARK_BLUE)

    add_styled_paragraph(doc, "迭代原则：", bold=True)
    add_bullet_list(doc, [
        "每个阶段结束时交付可运行的增量版本",
        "前端后端并行开发，通过 API 接口文档解耦",
        "第三方数据接入优先做 Mock 数据 → 再对接真实 API",
        "每周进行一次代码 Review 和进度对齐会议",
        "接口文档先行，前后端同步开发不互相阻塞",
    ])

    # 2.2
    add_heading_styled(doc, "2.2  人员任务分配", level=2)
    add_styled_paragraph(
        doc,
        "项目团队共 9 人，采用「前后端配对 Team + AI 专项组 + 运维」的组织模式。"
        "3 个 Team 各由 1 名前端 + 1 名后端组成，每个 Team 独立负责一组核心功能的全栈开发；"
        "AI 组专注大模型与 RAG 相关能力建设；运维负责部署、CI/CD 与测试。",
        first_line_indent=0.74,
    )

    create_table(doc,
        ["组别", "成员", "角色", "负责功能模块"],
        [
            ["Team A\n学业核心", "覃泽锴（前端）\n孔垂骄（后端）",
             "前端：课表/成绩/DDL 页面\n后端：教务系统对接、认证、课表/成绩/DDL 接口",
             "课表查询、成绩查询\nDDL 管理、选课推荐\n考试倒计时、一键评教"],
            ["Team B\n生活服务", "李亚飞（前端）\n谭旭睿（后端）",
             "前端：食堂/天气/快递/美食页面\n后端：第三方 API 对接、爬虫",
             "食堂导航、校外美食\n天气穿衣建议、快递追踪\n一卡通、图书馆"],
            ["Team C\n校园信息", "徐锐学（前端）\n毛立业（后端）",
             "前端：校车/校历/活动/通知页面\n后端：校园数据爬虫、通知聚合",
             "校车时刻、校历查询\n校园通知聚合、校园活动\n周末规划、Dashboard"],
            ["AI 组", "张傲楚\n谭博文",
             "LLM 意图路由与 Function Calling\nDeepTutor 集成与 RAG 问答\n智能出题与知识图谱",
             "AI 对话模块\n课件 RAG 问答\n智能出题、复习计划\n用户记忆、跨源推理"],
            ["运维", "朱圣相",
             "Docker 部署与容器编排\nCI/CD 流水线（GitHub Actions）\n集成测试与文档",
             "开发/生产环境搭建\n自动构建与部署\n全模块测试"],
        ],
        col_widths=[2.5, 3.5, 5, 4.5],
    )

    # 协作说明
    add_styled_paragraph(doc, "协作模式：", bold=True, space_before=6)
    add_bullet_list(doc, [
        "各 Team 内前后端紧密配合，接口文档先行，联调效率最大化",
        "AI 组为横向支撑，与三个 Team 协作提供 AI 能力（意图路由、RAG、推荐等）",
        "运维贯穿全程，各阶段末尾主导集成测试与部署验证",
        "每周全员例会对齐进度，双周迭代交付可运行版本",
    ])

    # 2.3
    add_heading_styled(doc, "2.3  进度计划", level=2)

    # 阶段一
    add_heading_styled(doc, "阶段一：需求分析与框架搭建（第 1-3 周）", level=3)
    add_styled_paragraph(
        doc,
        "目标：完成需求调研、技术选型验证、前后端项目初始化、数据库设计，为后续开发奠定基础。",
        first_line_indent=0.74,
    )
    create_table(doc,
        ["任务", "负责人", "交付物"],
        [
            ["详细需求调研与用例编写", "全员", "需求规格说明书"],
            ["前端项目初始化（Next.js + shadcn/ui + Tailwind）", "Team A/B/C 前端", "可运行的前端脚手架"],
            ["后端项目初始化（FastAPI + SQLAlchemy + Alembic）", "Team A/B/C 后端", "可运行的后端脚手架"],
            ["数据库表结构设计与迁移脚本", "Team A 后端", "ER 图 + Alembic 迁移文件"],
            ["教务系统登录流程调研与原型验证", "Team A 后端", "CAS 登录技术验证报告"],
            ["API 接口文档编写（前后端约定）", "全员", "OpenAPI / Swagger 文档"],
            ["DeepTutor 源码调研与技术方案", "AI 组", "RAG 技术选型文档"],
            ["第三方 API 申请（和风天气、高德等）", "Team B 后端", "API Key 清单"],
            ["Docker Compose 开发环境搭建", "运维（朱圣相）", "一键启动脚本"],
            ["CI/CD 流水线搭建（GitHub Actions）", "运维（朱圣相）", "自动构建 + 测试流程"],
        ],
        col_widths=[6.5, 3, 6],
    )

    # 阶段二
    add_heading_styled(doc, "阶段二：核心功能开发（第 4-6 周）", level=3)
    add_styled_paragraph(
        doc,
        "目标：完成 P0 核心功能模块开发，实现教务系统对接、课表/成绩/DDL 页面、食堂导航、校车时刻。",
        first_line_indent=0.74,
    )
    create_table(doc,
        ["任务", "负责人", "交付物"],
        [
            ["教务系统 CAS 登录对接（验证码 + JWT）", "Team A", "认证 API + 登录页面"],
            ["课程表查询接口 + 周视图页面", "Team A", "课表后端 API + 前端周视图"],
            ["成绩查询接口 + 成绩页面", "Team A", "成绩 API + 学期筛选/GPA 计算"],
            ["DDL 管理 CRUD + DDL 看板页面", "Team A", "DDL 增删改查 + 优先级排序"],
            ["食堂导航页面（静态数据）", "Team B", "6 个食堂 + 窗口品类 + 营业状态"],
            ["校车时刻页面", "Team C", "三校区线路 + 工作日/周末 + 倒计时"],
            ["校历页面（时间线视图）", "Team C", "学期事件 + 颜色区分"],
            ["Dashboard 仪表盘页面", "Team C", "今日课程 + 成绩汇总 + 学分进度"],
            ["DDL 微信/QQ/学习通数据同步调研", "Team A 后端", "技术可行性报告"],
            ["核心模块集成测试", "运维（朱圣相）", "测试用例 + 测试报告"],
        ],
        col_widths=[6.5, 3, 6],
    )

    # 阶段三
    add_heading_styled(doc, "阶段三：AI 接入与功能完善（第 7-9 周）", level=3)
    add_styled_paragraph(
        doc,
        "目标：接入 AI 大模型实现智能对话、Function Calling 意图路由、DeepTutor RAG 问答，"
        "完善天气/快递/美食等生活服务模块。",
        first_line_indent=0.74,
    )
    create_table(doc,
        ["任务", "负责人", "交付物"],
        [
            ["通义千问对话接入 + 聊天页面", "AI 组 + Team A 前端", "AI 对话 API + 聊天 UI"],
            ["Function Calling 意图路由（课表/成绩/DDL）", "AI 组", "5+ 意图自动识别与执行"],
            ["课件上传 + 向量化 + RAG 问答（DeepTutor）", "AI 组", "文件上传 → Embedding → 检索 → 生成"],
            ["AI 智能出题 + 知识图谱可视化", "AI 组 + Team A 前端", "练习题生成 + 知识图谱前端"],
            ["天气服务接入（和风天气 API）", "Team B", "实时天气 + 穿衣建议"],
            ["校外美食推荐（高德 POI + 美团/饿了么）", "Team B", "美食搜索 + 地图展示"],
            ["快递追踪接入（快递 100 / 菜鸟）", "Team B", "物流 API 聚合 + 到站提醒"],
            ["校园通知爬虫（教务处/学院）", "Team C", "多源爬虫 + AI 摘要"],
            ["深色模式 + 设置页面", "Team C 前端", "next-themes + 用户偏好"],
            ["选课推荐引擎（培养方案差异分析）", "Team A + AI 组", "推荐算法 + 前端推荐页"],
            ["用户记忆系统（偏好持久化）", "AI 组", "用户画像存储 + 检索"],
            ["功能完善阶段集成测试", "运维（朱圣相）", "测试报告"],
        ],
        col_widths=[6.5, 3, 6],
    )

    # 阶段四
    add_heading_styled(doc, "阶段四：测试优化与交付（第 10-12 周）", level=3)
    add_styled_paragraph(
        doc,
        "目标：全模块集成测试、性能优化、安全加固、Docker 部署、项目文档整理、答辩准备。",
        first_line_indent=0.74,
    )
    create_table(doc,
        ["任务", "负责人", "交付物"],
        [
            ["端到端集成测试（全模块回归）", "运维（朱圣相）", "测试用例 + 缺陷报告"],
            ["性能优化（缓存策略、SQL 优化、前端懒加载）", "Team A/B/C", "性能报告"],
            ["安全加固（SQL 注入、XSS、CSRF、API 鉴权）", "Team A/B/C 后端", "安全审计报告"],
            ["Morning Briefing 晨间推送", "AI 组 + Team C", "定时任务 + 多源数据聚合"],
            ["跨源推理引擎（多工具联合调用）", "AI 组", "综合建议生成"],
            ["用户体验走查 + Bug 修复", "全员", "Bug 列表 + 修复记录"],
            ["Docker 生产环境部署", "运维（朱圣相）", "部署文档 + 监控方案"],
            ["项目文档整理", "全员", "API 文档、用户手册、技术文档"],
            ["答辩 PPT 制作与排练", "全员", "答辩 PPT + 演示视频"],
        ],
        col_widths=[6.5, 3, 6],
    )

    # 预期成果
    add_styled_paragraph(doc, "预期成果：", bold=True, space_before=6)
    add_bullet_list(doc, [
        "可运行的 Web 应用（前后端完整功能）",
        "AI 智能对话能力（自然语言问答 + Function Calling + RAG 文档问答）",
        "核心功能模块完整（课表/成绩/DDL/食堂/校车/天气/快递等）",
        "第三方数据接入（微信/学习通 DDL 同步、美团/饿了么、快递 100、和风天气等）",
        "Docker 可部署的生产级应用",
        "完整的项目文档（需求文档、API 文档、部署文档、用户手册）",
    ])

    # 2.4 甘特图
    add_heading_styled(doc, "2.4  进度计划（甘特图）", level=2)
    add_styled_paragraph(doc, "开发阶段进度总览", size=11, bold=True)
    create_gantt_table(doc)

    add_styled_paragraph(doc, "各小组并行工作分布", size=11, bold=True)
    create_role_gantt_table(doc)

    add_styled_paragraph(
        doc,
        "说明：M1 = 框架搭建完成 | M2 = 核心功能可用 | M3 = AI 与生活服务上线 | M4 = 正式交付",
        size=9, color=RGBColor(0x88, 0x88, 0x88),
    )

    # ===================================================================
    # 3 支持条件
    # ===================================================================
    add_heading_styled(doc, "3  支持条件", level=1)

    # 3.1
    add_heading_styled(doc, "3.1  软硬件环境", level=2)

    add_heading_styled(doc, "开发环境", level=3)
    create_table(doc,
        ["类别", "配置要求"],
        [
            ["操作系统", "Windows 11 / macOS / Linux（团队成员各异）"],
            ["IDE", "VS Code + 插件（ESLint、Prettier、Python、Tailwind IntelliSense）"],
            ["Node.js", "v20 LTS+"],
            ["Python", "3.11+"],
            ["Docker", "Docker Desktop 4.x（含 Docker Compose）"],
            ["Git", "2.40+，GitHub 托管"],
            ["浏览器", "Chrome / Edge 最新版（开发调试用）"],
        ],
        col_widths=[3.5, 13],
    )

    add_heading_styled(doc, "生产环境", level=3)
    create_table(doc,
        ["类别", "配置", "说明"],
        [
            ["云服务器", "4 核 8GB+", "运行 FastAPI + Next.js + PostgreSQL + Redis"],
            ["数据库", "PostgreSQL 16", "用户数据、DDL、学业缓存等持久化存储"],
            ["缓存", "Redis 7", "会话管理、速率限制、热点数据缓存"],
            ["向量数据库", "Milvus / Qdrant", "课件 RAG 向量存储（DeepTutor 集成）"],
            ["对象存储", "MinIO / 阿里云 OSS", "课件文件（PPT/PDF）上传存储"],
            ["域名 + SSL", "HTTPS 证书", "生产环境安全访问"],
            ["容器编排", "Docker Compose", "多服务编排与一键部署"],
        ],
        col_widths=[3, 4, 9],
    )

    # 3.2
    add_heading_styled(doc, "3.2  技术支持", level=2)

    add_heading_styled(doc, "系统架构", level=3)
    add_styled_paragraph(
        doc,
        "系统采用前后端分离的三层架构：前端展示层（Next.js + React）通过 REST API / SSE "
        "与后端服务层（FastAPI）通信；后端服务层负责业务逻辑处理、第三方 API 调用和 AI 模型调度；"
        "数据存储层包含 PostgreSQL（关系数据）、Redis（缓存）、Milvus（向量数据）。",
        first_line_indent=0.74,
    )

    add_heading_styled(doc, "核心技术栈", level=3)
    create_table(doc,
        ["层次", "技术", "版本 / 说明"],
        [
            ["前端框架", "Next.js + React + TypeScript", "最新稳定版"],
            ["UI 组件", "shadcn/ui + Tailwind CSS", "最新"],
            ["状态管理", "Zustand", "5.x"],
            ["数据请求", "TanStack React Query + Axios", "5.x"],
            ["图表可视化", "ECharts / D3.js", "知识图谱可视化"],
            ["后端框架", "FastAPI + Uvicorn", "0.115+"],
            ["ORM", "SQLAlchemy 2.0（async）", "异步数据库访问"],
            ["任务队列", "Celery + Redis Broker", "异步任务处理"],
            ["向量检索", "LangChain + Milvus/Qdrant", "RAG 文档检索"],
            ["文档解析", "Unstructured / PyMuPDF", "PPT/PDF 解析与分块"],
            ["Embedding", "text-embedding-3-small / BGE-M3", "文档向量化"],
            ["测试", "Pytest（后端）+ Vitest（前端）", "单元测试 + 集成测试"],
            ["CI/CD", "GitHub Actions", "自动构建 + 测试 + 部署"],
            ["部署", "Docker Compose", "多服务容器编排"],
        ],
        col_widths=[3, 7, 6],
    )

    # 第三方服务
    add_heading_styled(doc, "第三方服务依赖", level=3)
    create_table(doc,
        ["服务", "用途", "接口方式"],
        [
            ["四川大学教务系统", "课表、成绩、考试安排、评教", "HTTP 爬虫（CAS 模拟登录）"],
            ["通义千问 / DeepSeek API", "LLM 对话、Function Calling、复习计划生成", "REST API"],
            ["和风天气 API", "实时天气 + 7 天预报 + 生活指数", "REST API（免费 1000 次/天）"],
            ["高德地图 API", "POI 搜索、路线规划", "REST API（免费配额）"],
            ["快递 100 API", "物流跟踪聚合（支持主流快递公司）", "REST API"],
            ["美团/饿了么", "校外餐饮商户数据", "开放 API / 数据抓取"],
            ["微信开放平台", "消息推送、DDL 同步", "OAuth + WebHook"],
            ["学习通", "课程 DDL 与作业同步", "模拟登录数据抓取"],
            ["OpenAI Embedding / BGE", "文档向量化（RAG 用）", "REST API / 本地模型"],
            ["四川大学图书馆系统", "图书查询、借阅记录", "HTTP 爬虫"],
        ],
        col_widths=[4, 5.5, 6.5],
    )

    # 关键技术方案
    add_heading_styled(doc, "关键技术方案", level=3)

    add_numbered_item(doc, 1, "DeepTutor 集成方案", [
        "复用其双循环推理架构（Plan → Retrieve → Reason → Verify）",
        "文档上传 → PyMuPDF/Unstructured 解析 → 分块 → BGE-M3 Embedding → Milvus 存储",
        "查询时 RAG 检索 Top-K → LLM 带引用生成回答",
        "知识图谱：从文档提取实体关系 → Neo4j/内存图 → ECharts 前端可视化",
    ])

    add_numbered_item(doc, 2, "DDL 多源同步方案", [
        "微信：通过微信开放平台 WebHook 接收消息，NLP 提取时间 + 任务",
        "学习通：模拟登录抓取课程作业列表，定时同步",
        "QQ：通过 QQ 机器人 API（如 go-cqhttp）接收消息",
        "统一写入 Deadline 表，去重 + 冲突检测",
    ])

    add_numbered_item(doc, 3, "外卖/快递数据接入方案", [
        "美团/饿了么：优先尝试开放 API，备选方案为数据抓取 + 缓存",
        "快递追踪：快递 100 聚合 API，支持主流快递公司",
        "京东/淘宝物流：用户授权后调用开放平台物流查询接口",
    ])

    add_numbered_item(doc, 4, "天气数据方案", [
        "和风天气 API：免费版支持 7 天预报 + 生活指数（穿衣/紫外线/雨伞）",
        "成都地区精准定位（望江/江安/华西三校区经纬度）",
        "结合 LLM 生成自然语言穿衣建议",
    ])

    # ===================================================================
    # Footer
    # ===================================================================
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.space_before = Pt(24)
    pPr = footer_p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="6" w:color="999999"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    run = footer_p.add_run(
        "文档版本：v1.0  |  编写日期：2026-03-21  |  项目周期：12 周"
    )
    set_run_font(run, "微软雅黑", 9, False, RGBColor(0x99, 0x99, 0x99))

    # Save
    doc.save(OUTPUT_PATH)
    print(f"文档已生成: {OUTPUT_PATH}")


if __name__ == "__main__":
    generate_document()
