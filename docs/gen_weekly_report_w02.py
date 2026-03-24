"""
基于软件项目周报模板生成第2周周报（Word格式）
内容严格对照项目开发计划书阶段一（第1-3周）任务
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree
import os

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
out_path = os.path.join(SCRIPT_DIR, "..", "weekly-reports", "谭博文小组软件项目周报.docx")
os.makedirs(os.path.dirname(out_path), exist_ok=True)

doc = Document()

# ============================================================
# 页面设置
# ============================================================
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.17)
section.right_margin = Cm(3.17)

# ============================================================
# 默认字体设置（全局）
# ============================================================
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)


def add_heading_text(text, size=26, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                     font_name='黑体', space_after=Pt(6)):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return p


def add_info_line(label, value, size=14, font_name='宋体'):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.5
    run_label = p.add_run(label)
    run_label.bold = True
    run_label.font.size = Pt(size)
    run_label.font.name = font_name
    run_label._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run_val = p.add_run(value)
    run_val.font.size = Pt(size)
    run_val.font.name = font_name
    run_val._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return p


def add_section_title(text, size=14, font_name='黑体'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    pPr = p._element.get_or_add_pPr()
    pBdr = etree.SubElement(pPr, qn('w:pBdr'))
    bottom = etree.SubElement(pBdr, qn('w:bottom'))
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C41230')
    return p


def add_body_text(text, size=12, font_name='宋体', bold=False, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.bold = bold
    return p


def set_cell_font(cell, text, size=11, bold=False, font_name='宋体',
                  align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.bold = bold


def set_cell_shading(cell, color):
    shading = etree.SubElement(cell._element.get_or_add_tcPr(), qn('w:shd'))
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), color)


def make_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_font(cell, h, size=11, bold=True)
        set_cell_shading(cell, 'C41230')
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            align = WD_ALIGN_PARAGRAPH.LEFT if ci == len(headers) - 1 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_font(cell, val, size=11, align=align)
            if ri % 2 == 1:
                set_cell_shading(cell, 'FDF0F0')

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    return table


# ============================================================
# 文档内容
# ============================================================

add_heading_text("软件项目周报", size=26, font_name='黑体')

# 分隔线
p_line = doc.add_paragraph()
p_line.paragraph_format.space_after = Pt(4)
pPr = p_line._element.get_or_add_pPr()
pBdr = etree.SubElement(pPr, qn('w:pBdr'))
top = etree.SubElement(pBdr, qn('w:top'))
top.set(qn('w:val'), 'single')
top.set(qn('w:sz'), '12')
top.set(qn('w:space'), '1')
top.set(qn('w:color'), 'C41230')

# 基本信息
add_info_line("时间：", "2026年3月18日 - 2026年3月24日")
add_info_line("项目名称：", "四川大学智能校园助手（SCU Assistant）")
add_info_line("项目组长：", "谭博文")
add_info_line("周报期数：", "第 2 周")

# ============================================================
# 一、本周工作内容
# ============================================================
add_section_title("一、本周工作内容")

add_body_text(
    "本周为项目阶段一（需求分析与框架搭建）的第二周，团队在上周完成项目选题和技术路线规划的基础上，"
    "重点推进了前后端基础框架搭建、数据库表结构设计、教务系统登录调研以及开发环境配置等工作。"
)

# 1. 前后端项目框架搭建
add_body_text("1. 前后端项目框架搭建", bold=True, indent=False)
add_body_text(
    "完成了前端和后端项目的基础框架初始化工作。前端基于 Next.js + React + TypeScript "
    "搭建了项目脚手架，集成了 shadcn/ui 组件库和 TailwindCSS 样式框架，建立了基本的页面路由结构"
    "和布局组件（侧边栏、顶栏）。后端基于 FastAPI + SQLAlchemy + Alembic 搭建了服务框架，"
    "配置了项目目录结构、依赖管理和数据库迁移工具。"
)

make_table(
    headers=["模块", "技术栈", "完成内容"],
    rows=[
        ["前端脚手架", "Next.js 14 + React 19 + TypeScript", "项目初始化、路由结构、布局组件、shadcn/ui + TailwindCSS 集成"],
        ["后端脚手架", "FastAPI + SQLAlchemy + Alembic", "项目初始化、目录结构、依赖配置、数据库迁移工具接入"],
    ],
    col_widths=[3, 5, 6.5]
)

# 2. 数据库表结构设计
add_body_text("")
add_body_text("2. 数据库表结构设计", bold=True, indent=False)
add_body_text(
    "由 Team A 后端主导，完成了核心数据库表结构的初步设计，包括用户表（users）和"
    "DDL 截止日期表（deadlines）。使用 SQLAlchemy ORM 定义了数据模型，"
    "并通过 Alembic 生成了对应的数据库迁移脚本，为后续功能开发提供了数据层基础。"
)

# 3. 教务系统登录流程调研
add_body_text("3. 教务系统登录流程调研", bold=True, indent=False)
add_body_text(
    "Team A 后端对四川大学教务系统的 CAS 统一认证登录流程进行了深入调研。"
    "梳理了登录的关键步骤：获取验证码图片、密码双重 MD5 加密、Cookie 传递机制、"
    "以及 student.urpSoft.cn 教务系统的会话维持方式。输出了教务系统登录流程技术文档，"
    "为下周实现登录接口奠定了基础。"
)

# 4. DeepTutor 源码调研
add_body_text("4. DeepTutor 源码调研与 RAG 技术方案", bold=True, indent=False)
add_body_text(
    "AI 组对 HKUDS/DeepTutor 开源项目进行了源码阅读和技术调研，"
    "重点分析了其双循环推理架构（Plan → Retrieve → Reason → Verify）。"
    "结合项目实际需求，初步确定了 RAG 技术方案：采用 ChromaDB 作为向量数据库，"
    "PyMuPDF 进行文档解析，通过 LLM API 生成 Embedding 和回答。"
)

# 5. API 接口文档编写
add_body_text("5. API 接口文档与前后端约定", bold=True, indent=False)
add_body_text(
    "团队初步编写了前后端 API 接口约定，明确了认证接口（登录/登出/Token 刷新）、"
    "学业数据接口（课表/成绩查询）和 DDL 管理接口（CRUD）的请求/响应格式，"
    "使前后端组能够并行开发而不互相阻塞。"
)

# 6. 开发环境配置
add_body_text("6. 开发环境配置与工具学习", bold=True, indent=False)
add_body_text(
    "运维同学完成了 Docker Compose 开发环境的初步搭建，编写了 PostgreSQL + Redis 的容器编排文件。"
    "同时，各组成员继续深入学习各自负责方向的技术栈，前端组进行了 Next.js 和 React "
    "的实战练习，后端组完成了 FastAPI 的官方教程学习。"
)

# ============================================================
# 二、本周未完成的工作及原因
# ============================================================
add_section_title("二、本周未完成的工作及原因")

make_table(
    headers=["序号", "未完成事项", "原因说明"],
    rows=[
        ["1", "CI/CD 流水线搭建", "GitHub Actions 配置需要结合具体的测试框架，计划在下周前后端代码更完善后再配置"],
        ["2", "第三方 API 申请", "和风天气、高德地图等 API Key 申请需要实名认证，部分审核尚未通过"],
        ["3", "详细需求规格说明书", "需求调研仍在进行中，用例编写需结合教务系统实际接口情况，预计下周完成"],
    ],
    col_widths=[1.5, 4.5, 8.5]
)

# ============================================================
# 三、下周工作计划
# ============================================================
add_section_title("三、下周工作计划")

make_table(
    headers=["序号", "计划事项", "负责人", "预期产出"],
    rows=[
        ["1", "教务系统 CAS 登录原型验证", "Team A 后端", "完成登录流程代码编写，验证能否成功获取教务数据"],
        ["2", "完善数据库表结构设计", "Team A 后端", "补充学业缓存表（academic_cache）等，输出 ER 图"],
        ["3", "完成需求规格说明书", "全员", "详细用例文档，涵盖 5 大模块核心功能"],
        ["4", "CI/CD 流水线搭建", "运维（朱圣相）", "GitHub Actions 自动构建 + 测试流程"],
        ["5", "第三方 API Key 申请完成", "Team B 后端", "和风天气、高德地图 API Key 到位"],
        ["6", "前后端基础页面原型开发", "前端组", "登录页面、首页 Dashboard 基础框架"],
    ],
    col_widths=[1.2, 4.8, 2.5, 6]
)

# ============================================================
# 四、项目整体进度
# ============================================================
add_section_title("四、项目整体进度")

add_body_text(
    "当前项目整体进度约15%，处于阶段一（需求分析与框架搭建）的中期。"
    "前后端项目框架已初始化完成，数据库核心表结构已设计，教务系统登录流程已调研清楚，"
    "团队技术栈学习基本到位。预计下周完成阶段一剩余任务，为第4周进入核心功能开发阶段做好准备。"
)

add_body_text("12周开发计划概览：", bold=True, indent=False)

make_table(
    headers=["阶段", "周次", "主要工作", "状态"],
    rows=[
        ["需求分析与框架搭建", "第1-3周", "需求调研、技术选型、项目初始化、数据库设计", "进行中 ▶"],
        ["核心功能开发", "第4-6周", "课表/成绩、食堂导航、校车时刻、DDL管理、教务对接", "待开始"],
        ["AI接入与功能完善", "第7-9周", "通义千问对话接入、Function Calling、Dashboard", "待开始"],
        ["测试优化与交付", "第10-12周", "性能优化、安全加固、用户测试、Docker部署、答辩", "待开始"],
    ],
    col_widths=[3.8, 2.2, 6.5, 2]
)

# ============================================================
# 保存
# ============================================================
doc.save(out_path)
print(f"✓ 已生成周报: {out_path}")
