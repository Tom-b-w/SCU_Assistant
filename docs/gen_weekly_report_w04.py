from __future__ import annotations

import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(SCRIPT_DIR, "..", "weekly-reports", "谭博文小组软件项目周报_4.docx")
os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)


def set_default_style(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)


def add_text_run(paragraph, text: str, *, size: int = 12, bold: bool = False, font: str = "宋体"):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = font
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    return run


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    add_text_run(p, text, size=26, bold=True, font="黑体")

    line = doc.add_paragraph()
    p_pr = line._element.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "12")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), "C41230")
    border.append(top)
    p_pr.append(border)


def add_info_line(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.5
    add_text_run(p, label, size=14, bold=True)
    add_text_run(p, value, size=14)


def add_section_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    add_text_run(p, text, size=14, bold=True, font="黑体")

    p_pr = p._element.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C41230")
    border.append(bottom)
    p_pr.append(border)


def add_body(doc: Document, text: str, *, bold: bool = False, indent: bool = True) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    add_text_run(p, text, size=12, bold=bold)


def set_cell_text(cell, text: str, *, size: int = 11, bold: bool = False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.25
    run = add_text_run(p, text, size=size, bold=bold)
    return run


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        run = set_cell_text(cell, header, size=11, bold=True)
        shade_cell(cell, "C41230")
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if c_idx == len(headers) - 1 else WD_ALIGN_PARAGRAPH.CENTER
            cell = table.rows[r_idx + 1].cells[c_idx]
            set_cell_text(cell, value, size=11, align=align)
            if r_idx % 2 == 1:
                shade_cell(cell, "FDF0F0")

    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)


def build_report() -> None:
    doc = Document()
    set_default_style(doc)

    add_title(doc, "软件项目周报")
    add_info_line(doc, "时间：", "2026年4月1日 - 2026年4月7日")
    add_info_line(doc, "项目名称：", "四川大学智能校园助手（SCU Assistant）")
    add_info_line(doc, "项目组长：", "谭博文")
    add_info_line(doc, "周报期数：", "第 4 周")

    add_section_title(doc, "一、本周工作内容")
    add_body(
        doc,
        "本周项目在完成《需求规格说明书》定稿的基础上，开始由需求分析阶段过渡到原型与基础开发阶段。"
        "团队一方面围绕最终版需求文档完成需求基线确认、模块拆分和任务细化，另一方面已在仓库中搭建前后端初步框架，"
        "形成了可展示的简单网页架构和核心服务骨架，为后续联调和功能迭代打下基础。"
    )

    add_body(doc, "1. 需求规格说明书整理与基线确认", bold=True, indent=False)
    add_body(
        doc,
        "结合最终版需求规格说明书，团队再次核对了系统范围、角色定义、功能边界和交付优先级，确认系统围绕学业管理、校园餐饮、"
        "校园基础服务、AI Agent 核心能力、用户认证五大模块展开，累计覆盖 22 条用户故事，并以 P0、P1、P2 三级优先级作为开发排期依据。"
    )
    add_table(
        doc,
        ["需求基线项", "本周确认结果"],
        [
            ["功能模块", "共 5 大模块：学业、餐饮、校园基础、AI Agent、用户认证"],
            ["核心用户故事", "共 22 条，已完成优先级划分并作为迭代输入"],
            ["分析模型", "用例图、活动图、顺序图、ER 图、类图共 8 张 UML 图表"],
            ["外部接口", "梳理教务系统、学习通、通义千问 API、和风天气 API 共 4 类接口"],
        ],
        widths=[4, 10.5],
    )

    add_body(doc, "")
    add_body(doc, "2. 简单网页架构与页面骨架初步搭建", bold=True, indent=False)
    add_body(
        doc,
        "前端已基于 Next.js App Router 启动页面结构搭建，形成“认证区 + 主功能区”的双布局模式。"
        "当前仓库中已落地登录页、主页仪表盘、AI 对话、课表、成绩、考试、DDL、RAG、通知、天气、校车、校历、食堂、设置等页面路由，"
        "并配套侧边栏、顶部栏、移动端导航等通用布局组件，能够支撑后续逐页补充真实业务逻辑。"
    )
    add_table(
        doc,
        ["页面层级", "当前进展"],
        [
            ["认证入口", "已完成账号密码登录页与学习通扫码登录页逻辑骨架"],
            ["主界面布局", "已完成 Sidebar、Topbar、MobileNav 等通用框架组件"],
            ["核心业务页面", "已建立 dashboard、chat、academic、campus、food、weather、notification、settings 等页面路由"],
            ["交互基础", "前端已封装 auth、academic、chat、deadline、weather、notification 等 API 调用模块"],
        ],
        widths=[3.5, 11],
    )

    add_body(doc, "")
    add_body(doc, "3. 后端服务骨架与数据库初步落地", bold=True, indent=False)
    add_body(
        doc,
        "后端已基于 FastAPI 建立 API Gateway 入口，完成认证、学业、聊天、DDL、RAG、天气、通知、学习通、记忆等多个服务路由的初步组织，"
        "并补充健康检查、跨域配置、异常处理等基础能力。数据库方面已通过 Alembic 建立用户、DDL、考试、通知、学习通会话、RAG 与记忆等核心表结构，"
        "与需求规格说明书中的 ER 图和类图保持一致。"
    )
    add_table(
        doc,
        ["后端部分", "本周进展"],
        [
            ["服务入口", "已完成 FastAPI 主应用、路由注册与 /health 健康检查"],
            ["业务模块", "已建立 auth、academic、chat、deadline、rag、weather、notification、chaoxing、memory 等服务目录"],
            ["数据层", "已配置 SQLAlchemy + Alembic，形成多张核心业务表迁移记录"],
            ["基础保障", "已具备 Redis 缓存接入、异常处理、中间件和部分单元测试基础"],
        ],
        widths=[3.5, 11],
    )

    add_body(doc, "")
    add_body(doc, "4. 需求到实现的对应关系初步建立", bold=True, indent=False)
    add_body(
        doc,
        "本周已开始将需求文档中的核心场景映射到实际实现路径。例如，用户认证模块已对应登录页、验证码接口和扫码登录流程；"
        "学业模块已对应课表、成绩、考试、DDL 等页面和服务；AI Agent 模块已预留聊天、RAG、长期记忆等能力入口。"
        "虽然多数功能仍处于原型或骨架状态，但页面分层、接口分层和数据模型分层已具备初步可追踪性。"
    )

    add_section_title(doc, "二、本周未完成的工作及原因")
    add_body(
        doc,
        "部分功能尚未进入完整联调阶段，主要原因是本周工作重心仍然放在需求说明书定稿与总体架构落地，开发资源优先用于搭建通用框架与页面骨架。"
    )
    add_body(doc, "1. 教务系统、学习通等外部接口尚未完成稳定联调，真实数据抓取流程仍需继续验证。", indent=False)
    add_body(doc, "2. AI 对话的 Function Calling、RAG 检索和记忆能力已具备模块入口，但完整业务闭环尚未全部跑通。", indent=False)
    add_body(doc, "3. 前端多数页面目前以结构展示和接口占位为主，细节样式优化、异常态处理与数据状态管理仍需继续完善。", indent=False)

    add_section_title(doc, "三、下周工作计划")
    add_body(doc, "1. 以需求规格说明书为基线，优先推进 P0 功能开发，重点完成认证、课表、DDL、聊天等核心流程联调。", indent=False)
    add_body(doc, "2. 继续完善前端页面交互与组件复用，补齐主页仪表盘、学业模块和 AI 对话页面的真实数据展示。", indent=False)
    add_body(doc, "3. 推进数据库表结构、缓存策略和接口异常处理细节，提升服务稳定性和可测试性。", indent=False)
    add_body(doc, "4. 对接外部系统接口并补充测试用例，逐步把原型页面转化为可演示的端到端功能。", indent=False)

    add_section_title(doc, "四、项目整体进度")
    add_body(
        doc,
        "截至 2026年4月7日，项目整体进度预计约为 35%。需求分析阶段已基本完成，正式需求基线已经形成；"
        "项目现已进入“基础框架搭建 + 核心功能启动开发”的过渡阶段。前端已经具备可展示的简单网页架构，"
        "后端也完成了多服务路由与核心数据表的初步建设，说明项目已从文档驱动顺利切换到工程实现阶段。"
    )
    add_table(
        doc,
        ["阶段", "当前状态"],
        [
            ["阶段一：需求分析", "已完成，需求规格说明书_final.docx 已作为后续开发基线"],
            ["阶段二：框架与原型搭建", "进行中，前后端基础骨架已落地"],
            ["阶段三：核心功能开发", "已启动，认证、学业、AI、通知等模块开始逐步实现"],
            ["阶段四：联调与优化", "尚未全面展开，待核心流程稳定后推进"],
        ],
        widths=[4.2, 10.3],
    )

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build_report()
