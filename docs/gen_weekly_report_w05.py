from __future__ import annotations

import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(SCRIPT_DIR, "..", "weekly-reports", "谭博文小组软件项目周报_5.docx")
os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)


def set_default_style(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)


def add_text_run(paragraph, text: str, *, size: int = 12, bold: bool = False, font: str = "宋体"):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = font
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    return run


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    add_text_run(p, text, size=26, bold=True, font="黑体")

    line = doc.add_paragraph()
    p_pr = line._element.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "12")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), "C41230")
    border.append(top)
    p_pr.append(border)


def add_info_line(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.5
    add_text_run(p, label, size=14, bold=True)
    add_text_run(p, value, size=14)


def add_section_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    add_text_run(p, text, size=14, bold=True, font="黑体")

    p_pr = p._element.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C41230")
    border.append(bottom)
    p_pr.append(border)


def add_body(doc: Document, text: str, *, bold: bool = False, indent: bool = True) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    add_text_run(p, text, size=12, bold=bold)


def set_cell_text(cell, text: str, *, size: int = 11, bold: bool = False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.25
    return add_text_run(p, text, size=size, bold=bold)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        run = set_cell_text(cell, header, size=11, bold=True)
        shade_cell(cell, "C41230")
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if c_idx == len(headers) - 1 else WD_ALIGN_PARAGRAPH.CENTER
            cell = table.rows[r_idx + 1].cells[c_idx]
            set_cell_text(cell, value, size=11, align=align)
            if r_idx % 2 == 1:
                shade_cell(cell, "FDF0F0")

    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)


def build_report() -> None:
    doc = Document()
    set_default_style(doc)

    add_title(doc, "软件项目周报")
    add_info_line(doc, "时间：", "2026年4月9日 - 2026年4月15日")
    add_info_line(doc, "项目名称：", "四川大学智能校园助手（SCU Assistant）")
    add_info_line(doc, "项目组长：", "谭博文")
    add_info_line(doc, "周报期数：", "第 5 周")

    add_section_title(doc, "一、本周工作内容")
    add_body(
        doc,
        "本周项目在需求规格说明书定稿与框架骨架基本落地的基础上，进一步推进系统设计阶段工作。"
        "团队围绕总体架构、模块边界、接口协议、数据结构、部署方案与异常处理机制进行了集中整理，"
        "形成了与当前代码仓库实现相对应的系统设计说明书，为下一阶段核心功能联调和端到端实现提供了统一设计基线。"
    )

    add_body(doc, "1. 系统设计说明书与设计基线整理", bold=True, indent=False)
    add_body(
        doc,
        "本周完成了《系统设计说明书》和《系统设计规格说明书》首版内容整理，"
        "将前期需求分析结果与当前前后端仓库结构进行对照，明确系统从“需求定义”转入“设计约束 + 实现映射”阶段。"
        "文档内容覆盖概要设计与详细设计两个层面，可直接作为后续开发、联调、测试和答辩展示的设计依据。"
    )
    add_table(
        doc,
        ["设计基线项", "本周完成情况"],
        [
            ["核心文档", "已形成《系统设计说明书_v1.0》与《系统设计规格说明书_v1.0》"],
            ["设计范围", "覆盖认证、学业、DDL、AI 对话、RAG、通知、天气、学习通、记忆等模块"],
            ["设计依据", "与需求规格说明书_final、docker-compose、前后端现有代码结构保持一致"],
            ["文档作用", "作为后续编码、联调、测试和验收的统一设计基线"],
        ],
        widths=[4, 10.5],
    )

    add_body(doc, "")
    add_body(doc, "2. 总体架构、模块划分与接口边界细化", bold=True, indent=False)
    add_body(
        doc,
        "围绕当前仓库实现，本周进一步细化了“前端展示层 - API Gateway 层 - 业务服务层 - 数据与外部依赖层”的四层架构，"
        "并将认证、学业、DDL、聊天、RAG、通知、天气、学习通、记忆等服务模块的职责边界进行统一说明。"
        "同时对 REST 与 SSE 两类接口形态、外部系统调用路径以及服务之间的依赖关系进行了梳理。"
    )
    add_table(
        doc,
        ["设计项", "当前进展"],
        [
            ["总体架构", "已明确前端、网关、业务服务、数据与外部依赖四层分工"],
            ["模块拆分", "已梳理 auth、academic、deadline、chat、rag、notification、weather、chaoxing、memory 等模块职责"],
            ["接口设计", "已统一 REST 接口与聊天 SSE 流式接口的使用场景和调用方式"],
            ["外部依赖", "已整理教务系统、学习通、LLM/Embedding、和风天气等接口接入关系"],
        ],
        widths=[3.5, 11],
    )

    add_body(doc, "")
    add_body(doc, "3. 数据结构、部署方案与异常处理设计补齐", bold=True, indent=False)
    add_body(
        doc,
        "本周同步整理了系统的数据设计与部署设计。数据层方面，对用户、学业缓存、DDL、通知、学习通会话、对话历史、用户记忆等核心表结构进行了归纳；"
        "部署层方面，明确了 Docker Compose 组织前端、后端、数据库与缓存服务的方式，以及开发环境使用 SQLite/内存缓存、部署环境切换 PostgreSQL/Redis 的策略。"
        "此外，还补充了统一错误响应、限流、降级和配置管理等横切设计。"
    )
    add_table(
        doc,
        ["设计维度", "本周成果"],
        [
            ["数据库设计", "已整理多张核心业务表、缓存表与会话表结构，并与 Alembic 迁移保持对应"],
            ["缓存与状态", "明确 Redis/内存缓存用于限流、会话、缓存预取和临时状态管理"],
            ["部署设计", "明确前端、网关、PostgreSQL、Redis 的容器化部署关系与环境变量配置"],
            ["容错机制", "补充统一异常处理、外部接口降级、限流与开发环境兜底策略"],
        ],
        widths=[3.5, 11],
    )

    add_body(doc, "")
    add_body(doc, "4. 设计文档与当前实现的映射关系建立", bold=True, indent=False)
    add_body(
        doc,
        "本周还对设计文档与实际代码目录之间的对应关系进行了梳理。前端页面路由、后端服务目录、共享基础设施、数据库迁移和配置文件等均已能在设计文档中找到明确映射，"
        "这意味着后续开发过程中可以围绕“文档设计 - 代码实现 - 联调测试”的路径持续推进，减少多人协作时的理解偏差。"
    )

    add_section_title(doc, "二、本周未完成的工作及原因")
    add_body(
        doc,
        "本周工作重点主要放在系统设计文档补齐和设计约束统一，因此部分面向演示的真实业务流程尚未完全打通。"
        "当前仍有若干模块停留在骨架实现或局部可用阶段，需要在下一周进入更深入的联调与验证。"
    )
    add_body(doc, "1. 教务系统、学习通等外部接口尚未完成全面稳定联调，真实数据同步流程仍需持续验证。", indent=False)
    add_body(doc, "2. AI 对话、RAG、长期记忆等模块虽已具备服务入口与设计说明，但端到端闭环能力仍需进一步打磨。", indent=False)
    add_body(doc, "3. 校车、校历、食堂等校园生活类页面目前以前端路由和展示骨架为主，后端服务与真实数据支撑尚未完全补齐。", indent=False)
    add_body(doc, "4. 自动化测试、部署验证和性能检查工作尚未系统展开，原因是本周资源优先投入到设计基线整理。", indent=False)

    add_section_title(doc, "三、下周工作计划")
    add_body(doc, "1. 以系统设计说明书为依据，优先推进认证、课表、成绩、DDL、聊天等 P0 核心流程的联调与演示闭环。", indent=False)
    add_body(doc, "2. 继续对接教务系统与学习通接口，提升缓存预取、DDL 同步和异常场景处理的稳定性。", indent=False)
    add_body(doc, "3. 完善 AI 对话的工具调用、上下文拼接、记忆抽取与流式输出体验，增强实际可用性。", indent=False)
    add_body(doc, "4. 补齐前端页面真实数据展示、空状态与错误态处理，并推进 Dashboard 聚合展示效果。", indent=False)
    add_body(doc, "5. 开始补充测试用例、部署验证和接口级联调记录，为后续答辩和演示做准备。", indent=False)

    add_section_title(doc, "四、项目整体进度")
    add_body(
        doc,
        "截至 2026年4月15日，项目整体进度预计约为 45%。目前需求分析阶段和系统设计阶段已基本完成，"
        "项目已进入“设计基线明确 + 核心功能实现推进”的阶段。前端已形成可展示的主界面和模块路由，"
        "后端已具备 API Gateway、核心业务服务、数据库迁移和部分真实接口能力，下一阶段重点将转向核心功能联调、外部接口稳定性验证与演示效果完善。"
    )
    add_table(
        doc,
        ["阶段", "当前状态"],
        [
            ["阶段一：需求分析", "已完成，需求规格说明书_final.docx 作为需求基线"],
            ["阶段二：系统设计", "已基本完成，系统设计说明书与设计规格说明书已形成首版"],
            ["阶段三：框架与核心功能实现", "进行中，前后端骨架与部分核心模块已落地"],
            ["阶段四：联调与测试", "尚未全面展开，待核心流程进一步打通后推进"],
            ["阶段五：演示优化与答辩准备", "未开始系统收尾，后续将结合联调结果逐步完善"],
        ],
        widths=[4.2, 10.3],
    )

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build_report()
