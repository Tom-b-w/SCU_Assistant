"""
基于软件项目周报模板生成第3周周报（Word格式）
本周为阶段一（需求分析与框架搭建）收尾周，重点完成需求规格说明书
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree
import os

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
out_path = os.path.join(SCRIPT_DIR, "..", "weekly-reports", "week03.docx")
os.makedirs(os.path.dirname(out_path), exist_ok=True)

doc = Document()

# ============================================================
# 页面设置
# ============================================================
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.17)
section.right_margin = Cm(3.17)

# ============================================================
# 默认字体设置（全局）
# ============================================================
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)


def add_heading_text(text, size=26, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                     font_name='黑体', space_after=Pt(6)):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return p


def add_info_line(label, value, size=14, font_name='宋体'):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.5
    run_label = p.add_run(label)
    run_label.bold = True
    run_label.font.size = Pt(size)
    run_label.font.name = font_name
    run_label._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run_val = p.add_run(value)
    run_val.font.size = Pt(size)
    run_val.font.name = font_name
    run_val._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return p


def add_section_title(text, size=14, font_name='黑体'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    pPr = p._element.get_or_add_pPr()
    pBdr = etree.SubElement(pPr, qn('w:pBdr'))
    bottom = etree.SubElement(pBdr, qn('w:bottom'))
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C41230')
    return p


def add_body_text(text, size=12, font_name='宋体', bold=False, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.bold = bold
    return p


def set_cell_font(cell, text, size=11, bold=False, font_name='宋体',
                  align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.bold = bold


def set_cell_shading(cell, color):
    shading = etree.SubElement(cell._element.get_or_add_tcPr(), qn('w:shd'))
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), color)


def make_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_font(cell, h, size=11, bold=True)
        set_cell_shading(cell, 'C41230')
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            align = WD_ALIGN_PARAGRAPH.LEFT if ci == len(headers) - 1 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_font(cell, val, size=11, align=align)
            if ri % 2 == 1:
                set_cell_shading(cell, 'FDF0F0')

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    return table


# ============================================================
# 文档内容
# ============================================================

add_heading_text("软件项目周报", size=26, font_name='黑体')

# 分隔线
p_line = doc.add_paragraph()
p_line.paragraph_format.space_after = Pt(4)
pPr = p_line._element.get_or_add_pPr()
pBdr = etree.SubElement(pPr, qn('w:pBdr'))
top = etree.SubElement(pBdr, qn('w:top'))
top.set(qn('w:val'), 'single')
top.set(qn('w:sz'), '12')
top.set(qn('w:space'), '1')
top.set(qn('w:color'), 'C41230')

# 基本信息
add_info_line("时间：", "2026年3月25日 - 2026年3月31日")
add_info_line("项目名称：", "四川大学智能校园助手（SCU Assistant）")
add_info_line("项目组长：", "谭博文")
add_info_line("周报期数：", "第 3 周")

# ============================================================
# 一、本周工作内容
# ============================================================
add_section_title("一、本周工作内容")

add_body_text(
    "本周为阶段一（需求分析与框架搭建）的收尾周，核心任务是完成《需求规格说明书》的编写与定稿。"
    "团队围绕需求梳理、用户故事编写、UML建模、需求评审等工作展开，"
    "最终输出了覆盖系统全部功能模块的正式需求基线文档。"
)

# 1. 用户故事与需求梳理
add_body_text("1. 用户故事编写与需求梳理", bold=True, indent=False)
add_body_text(
    "团队以'作为[角色]，我希望[功能]，以便[价值]'的标准格式，系统梳理并编写了22条用户故事，"
    "覆盖系统全部5大功能模块。通过与团队成员的多轮讨论，对每条用户故事进行了价值评估和优先级排定，"
    "最终按P0（MVP必须）、P1（重要增强）、P2（锦上添花）三级优先级进行了分层。"
)

make_table(
    headers=["模块", "用户故事数", "P0", "P1", "P2", "典型用户故事示例"],
    rows=[
        ["学业模块", "6", "4", "1", "1", "查看课表、查询成绩、DDL管理、学习通DDL同步"],
        ["餐饮模块", "3", "1", "1", "1", "查看食堂窗口、按口味偏好推荐、查看评价"],
        ["校园基础模块", "4", "2", "1", "1", "校车时刻查询、校历查看、天气查看、通知浏览"],
        ["AI Agent核心", "5", "3", "1", "1", "自然语言查课表、AI问答、文档问答(RAG)"],
        ["用户认证", "4", "3", "1", "0", "学号密码登录、学习通扫码登录、个人信息管理"],
    ],
    col_widths=[2.5, 1.8, 1, 1, 1, 7.2]
)

# 2. UML 建模
add_body_text("")
add_body_text("2. UML分析模型绘制", bold=True, indent=False)
add_body_text(
    "为确保需求描述的完整性和多维度覆盖，团队采用多种UML分析模型对系统进行了建模，"
    "共输出8张专业UML图表，嵌入需求规格说明书正文。"
)

make_table(
    headers=["图表类型", "数量", "覆盖范围"],
    rows=[
        ["用例图（Use-Case Diagram）", "1", "系统整体用例，标注Student与4个外部系统（JWC、学习通、通义千问API、和风天气API）的交互关系"],
        ["活动图（Activity Diagram）", "3", "用户登录流程（含密码/扫码双分支）、AI对话处理流程（LLM Function Calling路由+SSE流式输出）、DDL任务管理流程（手动+学习通自动同步）"],
        ["顺序图（Sequence Diagram）", "2", "教务系统CAS登录时序、AI自然语言对话时序（SSE流式传输）"],
        ["ER图（Entity-Relationship Diagram）", "1", "PostgreSQL数据库核心实体及关系，标注PK/FK和1:N关联"],
        ["类图（Class Diagram）", "1", "后端四层架构（API Gateway→AI服务层→业务服务层→数据访问层）核心类及依赖/调用关系"],
    ],
    col_widths=[3.5, 1.2, 9.8]
)

# 3. 用例规格描述
add_body_text("")
add_body_text("3. 核心用例规格描述", bold=True, indent=False)
add_body_text(
    "对22个核心用例逐一编写了详细的用例规格描述，包括用例名称、参与者、前置条件、"
    "基本事件流、备选事件流、后置条件等要素，确保每个功能点的行为定义清晰、无歧义。"
    "重点描述了教务系统CAS登录、AI对话Function Calling路由、学习通扫码登录等"
    "涉及多方交互的复杂用例。"
)

# 4. 非功能性需求
add_body_text("4. 非功能性需求定义", bold=True, indent=False)
add_body_text(
    "明确了系统在性能、安全、可用性、可维护性和兼容性五个维度的非功能性需求约束，"
    "为后续开发和测试提供了验收基准。"
)

make_table(
    headers=["维度", "关键要求"],
    rows=[
        ["性能需求", "普通API响应<500ms，AI对话首token<3s，系统支持100并发用户"],
        ["安全需求", "JWT Token认证，教务/学习通密码加密存储，HTTPS传输，XSS/CSRF防护"],
        ["可用性与可靠性", "系统可用率≥99%，外部接口降级容错，关键数据定期备份"],
        ["可维护性与可扩展性", "四层架构分离，模块化设计，API接口版本化，支持新模块热插拔"],
        ["兼容性与环境约束", "支持Chrome/Firefox/Safari/Edge主流浏览器，移动端响应式适配"],
    ],
    col_widths=[3.5, 11]
)

# 5. 外部接口需求
add_body_text("")
add_body_text("5. 外部接口需求梳理", bold=True, indent=False)
add_body_text(
    "梳理了系统与4个外部系统的接口依赖关系，明确了各接口的通信协议、数据格式和降级容错策略。"
    "包括教务系统（JWC）CAS统一认证接口、学习通移动端API、通义千问大模型API"
    "以及和风天气API，每个接口均需实现降级容错机制，确保单个外部服务不可用时不影响系统其他功能。"
)

# 6. 需求评审
add_body_text("6. 需求评审与文档定稿", bold=True, indent=False)
add_body_text(
    "组织团队全员对需求规格说明书进行了内部评审，逐条核对用户故事与用例描述的一致性，"
    "检查UML图表与文字描述的对应关系，确认优先级划分的合理性。评审通过后，"
    "形成了正式的需求基线文档（需求规格说明书_final.docx），作为后续设计、"
    "开发、测试和验收的统一依据。"
)

# ============================================================
# 二、本周未完成的工作及原因
# ============================================================
add_section_title("二、本周未完成的工作及原因")

make_table(
    headers=["序号", "未完成事项", "原因说明"],
    rows=[
        ["1", "接口规格详细定义（OpenAPI）", "用例规格已完成，但各API的详细请求/响应字段定义尚需结合实际开发逐步细化"],
        ["2", "需求可追溯性矩阵", "用户故事与用例已一一对应，但完整的需求追踪矩阵（用户故事→用例→设计→测试用例）计划在测试阶段补充"],
        ["3", "非功能需求量化验证方案", "性能、安全等非功能需求已定义指标，但具体的测试方案和工具选型待阶段四确定"],
    ],
    col_widths=[1.5, 4.5, 8.5]
)

# ============================================================
# 三、下周工作计划
# ============================================================
add_section_title("三、下周工作计划")

add_body_text(
    "阶段一需求分析工作已全部完成，下周团队将正式进入阶段二（核心功能开发），"
    "以需求规格说明书为基线，启动各模块的详细设计与编码实现。"
)

make_table(
    headers=["序号", "计划事项", "负责人", "预期产出"],
    rows=[
        ["1", "教务系统课表/成绩数据对接", "后端组", "基于需求文档用例，实现课表查询和成绩查询接口"],
        ["2", "食堂导航模块开发", "前端组 + 后端组", "食堂窗口数据录入，列表展示，按需求文档实现推荐功能"],
        ["3", "校车时刻查询功能", "后端组", "校车数据接入，按线路/时间段查询"],
        ["4", "DDL 管理功能完善", "前端组", "按需求文档完善DDL增删改查及学习通同步体验"],
        ["5", "AI对话模块Function Calling联调", "AI 组 + 后端组", "按需求文档AI模块用例，实现工具调用与流式输出"],
        ["6", "各模块API接口细化", "全员", "结合编码实际，补充完善OpenAPI接口规格"],
    ],
    col_widths=[1.2, 4.8, 2.8, 5.7]
)

# ============================================================
# 四、项目整体进度
# ============================================================
add_section_title("四、项目整体进度")

add_body_text(
    "当前项目整体进度约25%，阶段一（需求分析与框架搭建）的需求分析部分已全部完成。"
    "需求规格说明书已正式定稿，文档覆盖5大功能模块、22个核心用例，"
    "包含用户故事、用例图、活动图、顺序图、ER图、类图共8张UML图表，"
    "以及完整的非功能性需求和外部接口需求定义。"
    "该文档将作为后续阶段二至阶段四所有开发、测试和验收工作的需求基线。"
)

add_body_text("12周开发计划概览：", bold=True, indent=False)

make_table(
    headers=["阶段", "周次", "主要工作", "状态"],
    rows=[
        ["需求分析与框架搭建", "第1-3周", "需求调研、技术选型、项目初始化、数据库设计、需求规格说明书", "已完成 ✓"],
        ["核心功能开发", "第4-6周", "课表/成绩、食堂导航、校车时刻、DDL管理、教务对接", "即将开始 ▶"],
        ["AI接入与功能完善", "第7-9周", "通义千问对话接入、Function Calling、Dashboard", "待开始"],
        ["测试优化与交付", "第10-12周", "性能优化、安全加固、用户测试、Docker部署、答辩", "待开始"],
    ],
    col_widths=[3.8, 2.2, 6.5, 2]
)

add_body_text("")
add_body_text("需求规格说明书关键数据：", bold=True, indent=False)

make_table(
    headers=["指标", "数据"],
    rows=[
        ["功能模块", "5 大模块：学业、餐饮、校园基础、AI Agent核心、用户认证"],
        ["用户故事", "22 条，覆盖全部核心功能"],
        ["优先级划分", "P0（MVP）13项、P1（重要增强）5项、P2（锦上添花）4项"],
        ["UML 图表", "8 张：用例图1、活动图3、顺序图2、ER图1、类图1"],
        ["非功能需求维度", "5 个维度：性能、安全、可用性、可维护性、兼容性"],
        ["外部接口依赖", "4 个：教务系统JWC、学习通、通义千问API、和风天气API"],
    ],
    col_widths=[4, 10.5]
)

# ============================================================
# 保存
# ============================================================
doc.save(out_path)
print(f"✓ 已生成周报: {out_path}")
