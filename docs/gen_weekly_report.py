"""
基于软件项目周报模板生成第1周周报（Word格式）
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree
import os

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
out_path = os.path.join(SCRIPT_DIR, "..", "weekly-reports", "week01.docx")
os.makedirs(os.path.dirname(out_path), exist_ok=True)

doc = Document()

# ============================================================
# 页面设置
# ============================================================
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.17)
section.right_margin = Cm(3.17)

# ============================================================
# 默认字体设置（全局）
# ============================================================
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)


def add_heading_text(text, size=26, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                     font_name='黑体', space_after=Pt(6)):
    """添加标题文字"""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return p


def add_info_line(label, value, size=14, font_name='宋体'):
    """添加信息行（如：项目名称：xxx）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.5
    # 标签（加粗）
    run_label = p.add_run(label)
    run_label.bold = True
    run_label.font.size = Pt(size)
    run_label.font.name = font_name
    run_label._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    # 值
    run_val = p.add_run(value)
    run_val.font.size = Pt(size)
    run_val.font.name = font_name
    run_val._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return p


def add_section_title(text, size=14, font_name='黑体'):
    """添加分区标题（如：本周工作内容）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    # 添加下边框
    pPr = p._element.get_or_add_pPr()
    pBdr = etree.SubElement(pPr, qn('w:pBdr'))
    bottom = etree.SubElement(pBdr, qn('w:bottom'))
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C41230')
    return p


def add_body_text(text, size=12, font_name='宋体', bold=False, indent=True):
    """添加正文"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.bold = bold
    return p


def add_bullet(text, size=12, font_name='宋体'):
    """添加列表项"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return p


def set_cell_font(cell, text, size=11, bold=False, font_name='宋体',
                  align=WD_ALIGN_PARAGRAPH.CENTER):
    """设置表格单元格格式"""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.bold = bold


def set_cell_shading(cell, color):
    """设置单元格底色"""
    shading = etree.SubElement(cell._element.get_or_add_tcPr(),
                                qn('w:shd'))
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), color)


def make_table(headers, rows, col_widths=None):
    """创建格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_font(cell, h, size=11, bold=True)
        set_cell_shading(cell, 'C41230')
        # 表头白色字
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # 数据行
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            align = WD_ALIGN_PARAGRAPH.LEFT if ci == len(headers) - 1 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_font(cell, val, size=11, align=align)
            if ri % 2 == 1:
                set_cell_shading(cell, 'FDF0F0')

    # 列宽
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    return table


# ============================================================
# 文档内容
# ============================================================

# 标题
add_heading_text("软件项目周报", size=26, font_name='黑体')

# 分隔线
p_line = doc.add_paragraph()
p_line.paragraph_format.space_after = Pt(4)
pPr = p_line._element.get_or_add_pPr()
pBdr = etree.SubElement(pPr, qn('w:pBdr'))
top = etree.SubElement(pBdr, qn('w:top'))
top.set(qn('w:val'), 'single')
top.set(qn('w:sz'), '12')
top.set(qn('w:space'), '1')
top.set(qn('w:color'), 'C41230')

# 基本信息
add_info_line("时间：", "2026年3月11日 - 2026年3月17日")
add_info_line("项目名称：", "四川大学智能校园助手（SCU Assistant）")
add_info_line("项目组长：", "谭博文")
add_info_line("周报期数：", "第 1 周")

# ============================================================
# 一、本周工作内容
# ============================================================
add_section_title("一、本周工作内容")

add_body_text(
    "本周为项目正式启动的第一周，主要完成了项目选题确认、团队组建与分工明确、"
    "技术路线与技术栈选型、项目功能模块规划以及开题报告的准备与汇报工作。"
)

# 1. 项目选题与方案确定
add_body_text("1. 项目选题与方案确定", bold=True, indent=False)
add_body_text(
    "经团队讨论，确定项目选题为『四川大学智能校园助手』（SCU Assistant），"
    "旨在构建一个AI驱动的一站式智能校园信息服务平台。项目针对当前校园信息碎片化、"
    "教务系统体验差、缺乏智能化等痛点，为川大6万+师生提供课表查询、成绩查看、"
    "食堂导航、校车时刻、DDL管理和AI智能对话等核心功能。"
)

# 2. 团队人员分工
add_body_text("2. 团队人员分工", bold=True, indent=False)
add_body_text(
    "9人团队按专业方向划分为4个小组，分工如下："
)

make_table(
    headers=["小组", "成员", "主要职责"],
    rows=[
        ["前端组（2人）", "覃泽锴、李亚飞", "Next.js 页面开发、UI/UX 组件搭建"],
        ["后端组（3人）", "孔垂骄、谭旭睿、毛立业", "FastAPI 接口开发、数据库建模、教务系统对接"],
        ["AI 组（2人）", "张傲楚、谭博文", "LLM 意图路由、Function Calling 实现"],
        ["运维与管理（2人）", "朱圣相、徐锐学", "Docker 部署、项目管理、CI/CD"],
    ],
    col_widths=[3.5, 4, 7]
)

# 3. 技术路线确定
add_body_text("")  # 空行
add_body_text("3. 技术路线与技术栈选型", bold=True, indent=False)

make_table(
    headers=["层级", "技术栈"],
    rows=[
        ["前端展示层", "Next.js 14 + React 19 + shadcn/ui + TailwindCSS + Zustand"],
        ["后端服务层", "FastAPI + SQLAlchemy + Redis + JWT + Pydantic"],
        ["数据存储层", "PostgreSQL + Redis Cache"],
        ["AI 能力层", "通义千问 API（阿里云）+ Function Calling"],
        ["部署运维", "Docker + Docker Compose + GitHub Flow"],
    ],
    col_widths=[3.5, 11]
)

# 4. 数据来源确认
add_body_text("")
add_body_text("4. 项目数据来源确认", bold=True, indent=False)
add_body_text(
    "明确了项目四大数据来源：（1）教务系统——通过CAS统一认证登录获取课表、成绩、"
    "学分等数据；（2）校园公开信息——学校官网获取食堂窗口信息和校车时刻表；"
    "（3）用户自主录入——DDL截止日期、待办事项等；（4）AI大模型——通义千问API"
    "提供自然语言问答与意图识别能力。"
)

# 5. 开题报告
add_body_text("5. 开题报告准备", bold=True, indent=False)
add_body_text(
    "完成了开题报告PPT的制作，涵盖选题背景、需求分析、产品定位与功能规划、"
    "团队分工、技术路线和12周开发计划等内容，并建立了项目GitHub仓库"
    "（https://github.com/Tom-b-w/SCU_Assistant）。"
)

# ============================================================
# 二、本周未完成的工作及原因
# ============================================================
add_section_title("二、本周未完成的工作及原因")

make_table(
    headers=["序号", "未完成事项", "原因说明"],
    rows=[
        ["1", "前后端项目框架搭建", "本周重心在项目规划和开题准备，框架搭建安排在下周启动"],
        ["2", "数据库表结构设计", "需先完成需求细化，下周结合具体功能进行设计"],
        ["3", "教务系统登录流程调研", "CAS认证+验证码机制较复杂，需要更多时间研究"],
    ],
    col_widths=[1.5, 4.5, 8.5]
)

# ============================================================
# 三、下周工作内容
# ============================================================
add_section_title("三、下周工作计划")

make_table(
    headers=["序号", "计划事项", "负责人", "预期产出"],
    rows=[
        ["1", "学习 Next.js / React 前端技术", "前端组", "完成官方教程，能搭建基础页面"],
        ["2", "学习 FastAPI / Python 后端技术", "后端组", "完成官方教程，能编写基础API"],
        ["3", "学习 AI 大模型 API 调用方式", "AI 组", "跑通通义千问API调用示例"],
        ["4", "搭建前后端项目基础框架", "全体", "项目可本地运行，基本目录结构建立"],
        ["5", "数据库表结构初步设计", "后端组", "完成用户、课表、DDL等核心表设计"],
        ["6", "调研教务系统CAS登录流程", "后端组", "输出教务系统登录流程技术文档"],
    ],
    col_widths=[1.2, 4.8, 2.5, 6]
)

# ============================================================
# 四、项目整体进度
# ============================================================
add_section_title("四、项目整体进度")

add_body_text(
    "当前项目整体进度约10%，处于规划与学习阶段。已完成项目选题确认、团队分工、"
    "技术路线选型和功能模块规划，团队成员正在学习各自负责方向的技术栈，"
    "计划下周开始着手实际编码工作。"
)

# 12周计划表
add_body_text("12周开发计划概览：", bold=True, indent=False)

make_table(
    headers=["阶段", "周次", "主要工作", "状态"],
    rows=[
        ["需求分析与框架搭建", "第1-3周", "需求调研、技术选型、项目初始化、数据库设计", "进行中"],
        ["核心功能开发", "第4-6周", "课表/成绩、食堂导航、校车时刻、DDL管理、教务对接", "待开始"],
        ["AI接入与功能完善", "第7-9周", "通义千问对话接入、Function Calling、Dashboard", "待开始"],
        ["测试优化与交付", "第10-12周", "性能优化、安全加固、用户测试、Docker部署、答辩", "待开始"],
    ],
    col_widths=[3.8, 2.2, 6.5, 2]
)

# ============================================================
# 保存
# ============================================================
doc.save(out_path)
print(f"✓ 已生成周报: {out_path}")
