# -*- coding: utf-8 -*-
from __future__ import annotations

import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_DIR = os.path.join(SCRIPT_DIR, "..", "weekly-reports")
os.makedirs(OUT_DIR, exist_ok=True)


def set_default_style(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)


def add_text_run(paragraph, text: str, *, size: int = 12, bold: bool = False, font: str = "宋体"):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = font
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    return run


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    add_text_run(p, text, size=26, bold=True, font="黑体")

    line = doc.add_paragraph()
    p_pr = line._element.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "12")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), "C41230")
    border.append(top)
    p_pr.append(border)


def add_info_line(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.5
    add_text_run(p, label, size=14, bold=True)
    add_text_run(p, value, size=14)


def add_section_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    add_text_run(p, text, size=14, bold=True, font="黑体")

    p_pr = p._element.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C41230")
    border.append(bottom)
    p_pr.append(border)


def add_body(doc: Document, text: str, *, bold: bool = False, indent: bool = True) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    add_text_run(p, text, size=12, bold=bold)


def set_cell_text(cell, text: str, *, size: int = 11, bold: bool = False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.25
    return add_text_run(p, text, size=size, bold=bold)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        run = set_cell_text(cell, header, size=11, bold=True)
        shade_cell(cell, "C41230")
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if c_idx == len(headers) - 1 else WD_ALIGN_PARAGRAPH.CENTER
            cell = table.rows[r_idx + 1].cells[c_idx]
            set_cell_text(cell, value, size=11, align=align)
            if r_idx % 2 == 1:
                shade_cell(cell, "FDF0F0")

    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)


def build_report(cfg: dict) -> str:
    doc = Document()
    set_default_style(doc)

    add_title(doc, "软件项目周报")
    add_info_line(doc, "时间：", cfg["time"])
    add_info_line(doc, "项目名称：", "四川大学智能校园助手（SCU Assistant）")
    add_info_line(doc, "项目组长：", "谭博文")
    add_info_line(doc, "周报期数：", f"第 {cfg['week_no']} 周")

    add_section_title(doc, "一、本周工作内容")
    add_body(doc, cfg["intro"])

    for item in cfg["work_items"]:
        add_body(doc, item["title"], bold=True, indent=False)
        add_body(doc, item["text"])
        if item.get("table"):
            headers, rows, widths = item["table"]
            add_table(doc, headers, rows, widths)
        add_body(doc, "")

    add_section_title(doc, "二、本周未完成的工作及原因")
    add_body(doc, cfg["unfinished_intro"])
    for idx, line in enumerate(cfg["unfinished"], 1):
        add_body(doc, f"{idx}. {line}", indent=False)

    add_section_title(doc, "三、下周工作计划")
    for idx, line in enumerate(cfg["next_week"], 1):
        add_body(doc, f"{idx}. {line}", indent=False)

    add_section_title(doc, "四、项目整体进度")
    add_body(doc, cfg["progress_text"])
    headers, rows, widths = cfg["progress_table"]
    add_table(doc, headers, rows, widths)

    out_path = os.path.join(OUT_DIR, cfg["filename"])
    doc.save(out_path)
    return out_path


WEEK_6 = {
    "filename": "谭博文小组软件项目周报_6.docx",
    "week_no": 6,
    "time": "2026年4月16日 - 2026年4月22日",
    "intro": "本周项目在系统设计首版文档的基础上，继续推进设计深化与基线收敛工作。团队围绕《系统设计说明书最终版》的定稿，集中补齐了总体架构、接口协议、数据库设计、部署方案、异常处理和详细模块设计等内容，使系统从“有代码骨架”进一步过渡到“有完整设计约束和实现映射”的阶段。",
    "work_items": [
        {
            "title": "1. 《系统设计说明书最终版》定稿与设计基线收敛",
            "text": "本周完成了《系统设计说明书最终版》的整理与定稿，进一步统一了项目背景、术语定义、设计目标和文档结构，将需求规格说明书中的功能边界与当前仓库中的前后端实现进行逐项对照，形成了可直接支撑后续联调、测试和答辩展示的正式设计基线。",
            "table": (
                ["设计章节", "本周完成情况"],
                [
                    ["引言与总体设计", "补充项目背景、术语定义、四层架构说明，明确展现层、业务层、持久化层和数据库层职责"],
                    ["接口设计", "整理 Web 前端、教务系统、学习通、天气、LLM/Embedding 等外部接口，并统一内部服务调用关系"],
                    ["数据设计", "归纳 10 张核心数据库表和 Redis 缓存 Key 结构，明确字段、约束和用途"],
                    ["部署与出错处理", "明确 Docker Compose 部署拓扑、环境变量配置和统一异常返回策略"],
                ],
                [4, 10.5],
            ),
        },
        {
            "title": "2. 详细设计规格补充与模块职责细化",
            "text": "在概要设计完成的基础上，团队继续补充详细设计内容，重点说明各模块内部结构、核心类/函数职责、关键业务流程和界面对应关系，使文档不仅描述“系统是什么”，也说明“系统如何实现”。",
            "table": (
                ["模块", "本周补充设计内容"],
                [
                    ["认证与教务", "补充验证码登录、JWT 刷新、教务会话缓存和课表/成绩缓存读取流程说明"],
                    ["AI 对话与 RAG", "细化 Claude Tool Use、知识库上传解析、向量检索和问答生成流程"],
                    ["DDL 与学习通", "明确 Deadline 表设计、基础 CRUD、学习通扫码绑定与 DDL 同步路径"],
                    ["天气/通知/简报", "补充天气查询、通知抓取、每日简报聚合逻辑以及页面入口说明"],
                ],
                [3.8, 10.7],
            ),
        },
        {
            "title": "3. 文档与当前实现映射关系校验",
            "text": "本周对设计文档与现有代码目录进行了对照校验。前端页面路由、状态管理、API 调用层，以及后端网关、业务服务、共享基础设施、数据库迁移等内容均已能够在文档中找到清晰映射，降低了后续多人协作开发中的理解偏差。",
            "table": (
                ["实现层", "当前映射结果"],
                [
                    ["前端页面层", "已映射 login、dashboard、chat、academic、weather、notification、settings 等页面与布局组件"],
                    ["前端状态与请求层", "已映射 auth-store、chat-store 及各业务 lib 调用封装"],
                    ["后端服务层", "已映射 academic、chat、rag、deadline、weather、notification、briefing、chaoxing、memory、quiz 等模块"],
                    ["共享基础设施层", "已映射 config、database、cache、exceptions、models、llm_client 等公共能力"],
                ],
                [3.8, 10.7],
            ),
        },
        {
            "title": "4. 面向演示的页面结构与业务流程整理",
            "text": "为了后续课程答辩和演示准备，团队同步梳理了系统主要页面与业务流程入口，明确了登录页、首页仪表盘、AI 对话页、课表页、DDL 管理页、RAG 问答页、天气页、通知页等核心页面的展示职责和交互路径，增强了系统的可展示性与演示连贯性。",
        },
    ],
    "unfinished_intro": "本周工作重点放在系统设计说明书最终版的补齐和设计基线统一，因此部分偏实现和联调性质的工作仍未完全展开，当前还存在若干需要在下一周继续推进的事项。",
    "unfinished": [
        "教务系统、学习通等外部接口尚未完成全面稳定联调，真实数据抓取与同步流程仍需持续验证。",
        "前端多数页面虽然已有完整结构和设计说明，但真实数据展示、空状态与错误态处理仍需进一步补齐。",
        "自动化测试、性能检查和部署验证工作尚未系统展开，原因是本周资源优先投入到设计文档定稿与详细设计整理。",
    ],
    "next_week": [
        "以《系统设计说明书最终版》为依据，优先推进认证、课表、成绩、DDL、聊天等 P0 核心流程的联调与闭环验证。",
        "继续对接教务系统与学习通接口，重点验证缓存预取、作业同步和异常场景处理的稳定性。",
        "完善 AI 对话、RAG、记忆和每日简报等模块的真实调用链，增强实际可用性与演示效果。",
        "补充部署验证和基础测试记录，为后续答辩材料和系统演示做准备。",
    ],
    "progress_text": "截至 2026年4月22日，项目整体进度预计约为 55%。目前需求分析阶段和系统设计阶段已基本完成，系统设计说明书最终版已经形成，项目已进入“设计基线稳定 + 核心流程联调准备”的阶段。前后端骨架、数据库模型、接口结构和主要页面信息架构均已明确，下一阶段工作重点将转向核心功能验证和端到端演示闭环。",
    "progress_table": (
        ["阶段", "当前状态"],
        [
            ["阶段一：需求分析", "已完成，需求规格说明书已作为后续开发基线"],
            ["阶段二：系统设计", "已完成最终版整理，总体设计与详细设计内容已收敛"],
            ["阶段三：框架与核心功能实现", "进行中，前后端骨架和多模块服务已落地"],
            ["阶段四：联调与测试", "即将重点推进，待核心流程逐步打通后全面展开"],
            ["阶段五：演示优化与答辩准备", "未系统收尾，后续将结合联调结果持续完善"],
        ],
        [4, 10.5],
    ),
}


WEEK_7 = {
    "filename": "谭博文小组软件项目周报_7.docx",
    "week_no": 7,
    "time": "2026年4月23日 - 2026年4月29日",
    "intro": "本周项目在《系统设计说明书最终版》定稿的基础上，继续推进“设计到实现”的衔接工作。团队围绕核心模块的实现路径、接口调用约束、缓存与异常处理策略、前端页面信息架构和演示链路进行了进一步梳理，为下一步核心流程联调、测试与答辩展示打下了更扎实的工程基础。",
    "work_items": [
        {
            "title": "1. 核心模块实现路径与接口约束再细化",
            "text": "本周对认证、学业、AI、DDL、学习通、校园资讯等核心模块的实现路径进行了再次收敛，重点核对接口入口、返回结构、依赖关系和调用顺序，确保各模块在联调时能够围绕统一的数据结构和错误语义推进。",
            "table": (
                ["联调对象", "本周落实情况"],
                [
                    ["认证链路", "梳理验证码获取、账号登录、JWT 刷新、当前用户信息查询等接口边界"],
                    ["学业链路", "统一课表、成绩、培养方案、考试等接口的缓存优先策略与刷新入口"],
                    ["AI 链路", "进一步明确 Chat、Tool Use、RAG、Memory、Quiz 等模块之间的数据流转关系"],
                    ["学习通与 DDL", "核对扫码绑定、会话存储、作业同步写入 Deadline 表的实现路径"],
                ],
                [3.8, 10.7],
            ),
        },
        {
            "title": "2. AI 能力与知识库演示闭环准备",
            "text": "围绕系统的 AI-First 特点，本周重点整理了对话、知识库问答、用户记忆和简报生成相关能力的展示路径，明确哪些功能适合在答辩中重点演示，哪些功能需要继续补齐端到端闭环和异常兜底。",
            "table": (
                ["能力点", "当前准备情况"],
                [
                    ["AI 对话", "已具备自然语言问答、流式输出和工具调用设计路径，可作为核心演示入口"],
                    ["RAG 课件问答", "已明确知识库创建、文档上传、向量检索、带引用回答的完整流程"],
                    ["用户记忆", "已补充偏好抽取与存储说明，为后续个性化回答增强预留能力"],
                    ["每日简报/智能出题", "已形成聚合数据生成简报、基于知识库生成题目的设计说明，后续继续打磨演示效果"],
                ],
                [3.8, 10.7],
            ),
        },
        {
            "title": "3. 数据、缓存与异常处理一致性整理",
            "text": "本周同步复核了数据库表结构、缓存 Key 设计、限流策略和异常返回格式，确保服务在真实外部依赖不稳定或开发环境资源受限时，仍能通过缓存、降级或统一错误响应保证系统可用性。",
            "table": (
                ["设计维度", "本周整理结果"],
                [
                    ["数据库结构", "继续核对 users、academic_cache、deadlines、documents、user_memories、chaoxing_sessions 等表之间的关系"],
                    ["缓存策略", "明确教务会话、刷新令牌、限流计数器和二维码状态等短生命周期数据的缓存使用方式"],
                    ["异常处理", "统一 SESSION_EXPIRED、RATE_LIMITED、INTERNAL_ERROR 等错误语义，便于前后端协同处理"],
                    ["开发兜底", "保持 SQLite、fakeredis、Mock 数据等轻量方案，降低联调和演示门槛"],
                ],
                [3.8, 10.7],
            ),
        },
        {
            "title": "4. 页面信息架构与答辩展示路径梳理",
            "text": "团队进一步整理了登录页、首页仪表盘、AI 对话、课表、成绩、DDL、RAG、天气、通知、校车、校历、食堂和设置等页面的展示顺序与导航关系，为后续形成连贯的系统演示路线和答辩讲解逻辑提供支持。",
        },
    ],
    "unfinished_intro": "本周主要工作仍然集中在实现路径统一、模块衔接和演示准备层面，因此部分需要真实数据支撑的功能闭环尚未完全完成，后续仍需继续补齐。",
    "unfinished": [
        "教务系统、学习通、通知抓取等外部依赖的稳定性验证仍不充分，部分真实场景下的超时、失效和重试逻辑需要继续打磨。",
        "Quiz、Memory、食堂、校车、校历等扩展能力已有页面或服务入口，但端到端可演示程度仍不均衡。",
        "自动化测试、部署演练和性能安全验证尚未形成完整闭环，原因是当前资源仍优先保障核心流程实现与答辩准备。",
    ],
    "next_week": [
        "集中推进教务登录、课表成绩读取、学习通同步、DDL 管理、AI 对话等核心链路的联调和演示闭环。",
        "补齐前端页面的真实数据展示、空状态、错误态和交互反馈，提升整体可用性与展示完成度。",
        "补充接口级测试、部署验证和典型异常场景验证记录，增强系统稳定性和可答辩性。",
        "围绕最终版设计文档整理答辩讲解材料，形成从需求、设计到实现的完整汇报链路。",
    ],
    "progress_text": "截至 2026年4月29日，项目整体进度预计约为 62%。目前项目已经完成需求分析和系统设计两大阶段，进入“核心流程联调准备 + 演示路径收敛”的关键阶段。前后端主干结构、主要服务模块、数据库与缓存设计、页面信息架构和异常处理策略均已明确，后续重点将转向真实功能闭环、测试验证和答辩展示优化。",
    "progress_table": (
        ["阶段", "当前状态"],
        [
            ["阶段一：需求分析", "已完成，需求规格说明书作为需求基线持续使用"],
            ["阶段二：系统设计", "已完成，系统设计说明书最终版已形成统一设计依据"],
            ["阶段三：核心功能实现", "进行中，核心模块实现路径和接口边界已进一步明确"],
            ["阶段四：联调与测试", "逐步启动，下一阶段将重点推进真实链路验证"],
            ["阶段五：演示优化与答辩准备", "已开始梳理展示路径，后续将结合联调结果持续完善"],
        ],
        [4, 10.5],
    ),
}


if __name__ == "__main__":
    for report in (WEEK_6, WEEK_7):
        print(build_report(report))
